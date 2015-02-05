__author__ = 'nicole'

from rdflib.namespace import FOAF, DC, RDFS, OWL

import tarfile

import re
import os

from models.Assoc import Assoc
from sources.Source import Source
from models.Dataset import Dataset
from docx import Document
import curie_map
from utils.GraphUtils import GraphUtils
import glob


class GeneReviews(Source):
    '''
    THIS DATA SOURCE IS IN DEVELOPMENT AND IS SUBJECT TO CHANGE
    A stub for processing GeneReviews Word documents.  This is testing on some local versions
    of word documents as a test, to produce disease ids (NBKids), together with labels and descriptions.

    '''

    files = {
    }

    def __init__(self,args=[]):
        Source.__init__(self, 'genereviews')

        self.load_bindings()

        self.dataset = Dataset('genereviews', 'Gene Reviews', 'http://genereviews.org/')

        #data-source specific warnings (will be removed when issues are cleared)

        return

    def fetch(self,is_dl_forced):
        '''
        No fetching method at this time; the data is not yet available remotely.
        :return: None
        '''

        #self.get_files(is_dl_forced)

        return

    def parse(self,limit=None):
        '''
        abstract method to parse all data from an external resource, that was fetched in
        fetch()
        this should be overridden by subclasses
        :return: None
        '''

        #loop through and parse each file in the raw dir
        i=0
        for k in glob.glob(self.rawdir+'/*.docx'):
            i += 1
            if (limit is not None and i >= limit):
                break

            #print('gonna parse',k,'!')
            self._get_data(k)
            #TODO add each file, and the version information on a file-level basis

        #k = ('/').join((self.rawdir,'Incontinentia_Pigmenti_GeneReview.docx'))
        #self._get_data(k)
        self.load_bindings()

        print("INFO: Found", len(self.graph), "nodes")


        return


    def _get_data(self,file):
        '''
        This will add a disease node to the graph from one GeneReviews file.
        If the genereviews disease id (NBK*) can't be found in the document, it will write an error
        and no triples will be added
        :param file:  the file to parse
        :return:
        '''
        line_counter = 0
        document = Document(file)
        par = document.paragraphs
        tables = document.tables
        print('num paragraphs',len(par))
        pnum = 0
        disease_name = None
        description = None
        disease_num = None
        synonyms = None
        #cu = CurieUtil(curie_map.get())
        gu = GraphUtils(curie_map.get())
        for p in par:
            pnum +=1

            if (p.style == 'Heading1'):
                disease_name = p.text.strip()
            elif (re.match('Synonym',p.text)):
                #TODO are these a list?
                synonyms = re.sub('Synonym\:?','',p.text).strip()
            elif (re.match('Disease characteristics',p.text)):
                description = re.sub('Disease characteristics\w*\.?','',p.text).strip()
            elif (re.search('NBK\d+',p.text)):
                m = re.search('(NBK\d+)',p.text).group(0)
                disease_num = m

        if (disease_num is None):
            print("ERROR: could not find id for",file)
        else:
            disease_id = 'GeneReviews:'+disease_num
            #TODO add equivalences or types
            gu.addClassToGraph(self.graph,disease_id,disease_name,None,description)
            if synonyms is not None:
                gu.addSynonym(self.graph,disease_id,synonyms,Assoc.relationships['hasExactSynonym'])


        return


