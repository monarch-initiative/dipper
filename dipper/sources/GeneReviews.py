__author__ = 'nicole'

import re
import glob
import csv
import logging

from docx import Document

from dipper.models.Assoc import Assoc
from dipper.sources.Source import Source
from dipper.models.Dataset import Dataset
from dipper import curie_map
from dipper.utils.GraphUtils import GraphUtils

logger = logging.getLogger(__name__)


class GeneReviews(Source):
    """
    WARNING: THIS DATA SOURCE IS IN DEVELOPMENT AND IS SUBJECT TO CHANGE
    A stub for processing GeneReviews Word documents.  This is testing on some local versions
    of word documents as a test, to produce disease ids (NBKids), together with labels and descriptions.
    """

    files = {
        'idmap': {'file': 'NBKid_shortname_OMIM.txt',
                  'url': 'http://ftp.ncbi.nih.gov/pub/GeneReviews/NBKid_shortname_OMIM.txt'},
        'titles': {'file': 'GRtitle_shortname_NBKid.txt',
                   'url': 'http://ftp.ncbi.nih.gov/pub/GeneReviews/GRtitle_shortname_NBKid.txt'}
        }

    def __init__(self):
        Source.__init__(self, 'genereviews')

        self.load_bindings()

        self.dataset = Dataset('genereviews', 'Gene Reviews', 'http://genereviews.org/',
                               None, 'http://www.ncbi.nlm.nih.gov/books/NBK138602/')

        # data-source specific warnings (will be removed when issues are cleared)

        return

    def fetch(self, is_dl_forced):
        """
        We fetch the general GeneReviews files from NCBI.
        Note that the doc files are not yet available remotely.
        :return: None
        """

        self.get_files(is_dl_forced)

        return

    def parse(self, limit=None):
        """
        :return: None
        """

        # loop through and parse each file in the raw dir
        i = 0
        for k in glob.glob(self.rawdir+'/*.docx'):
            i += 1
            if limit is not None and i >= limit:
                break

            logger.debug('gonna parse %s!', k)
            self._get_data(k)
            # TODO add each file, and the version information on a file-level basis

        # in development
        # k = ('/').join((self.rawdir,'Incontinentia_Pigmenti_GeneReview.docx'))
        # self._get_data(k)

        if self.testOnly:
            self.testMode = True

        self._get_equivids(limit)
        self._get_titles(limit)

        self.load_bindings()

        # no test subset for now; test == full graph
        self.testgraph = self.graph

        logger.info("Found %d nodes", len(self.graph))

        return

    def _get_data(self, file):
        """
        This will add a disease node to the graph from one GeneReviews file.
        If the genereviews disease id (NBK*) can't be found in the document, it will write an error
        and no triples will be added
        :param file:  the file to parse
        :return:
        """

        line_counter = 0
        document = Document(file)
        par = document.paragraphs
        tables = document.tables
        logger.info('num paragraphs %d', len(par))
        pnum = 0
        disease_name = None
        description = None
        disease_num = None
        synonyms = None
        gu = GraphUtils(curie_map.get())

        for p in par:
            pnum += 1

            if p.style == 'Heading1':
                disease_name = p.text.strip()
            elif re.match('Synonym', p.text):
                # TODO are these a list?
                synonyms = re.sub('Synonym\:?', '', p.text).strip()
            elif re.match('Disease characteristics',p.text):
                description = re.sub('Disease characteristics\w*\.?', '', p.text).strip()
            elif re.search('NBK\d+',p.text):
                m = re.search('(NBK\d+)',p.text).group(0)
                disease_num = m

        if disease_num is None:
            logger.error("could not find id for %s", file)
        else:
            disease_id = 'GeneReviews:'+disease_num
            # TODO add equivalences or types
            gu.addClassToGraph(self.graph, disease_id, disease_name, None, description)
            if synonyms is not None:
                gu.addSynonym(self.graph, disease_id, synonyms, Assoc.properties['hasExactSynonym'])

        return

    def _get_equivids(self, limit):
        """
        The file processed here is of the format:
        #NBK_id GR_shortname    OMIM
        NBK1103 trimethylaminuria       136132
        NBK1103 trimethylaminuria       602079
        NBK1104 cdls    122470
        Where each of the rows represents a mapping between a gr id and an omim id.
        These are a 1:many relationship, and some of the omim ids are genes (not diseases).
        Therefore, we need to create a loose coupling here.  We make the assumption that these NBKs are
        generally higher-level grouping classes; therefore the OMIM ids are treated as subclasses.  (This
        assumption is poor for those omims that are actually genes, but we have no way of knowing what those
        are here...we will just have to deal with that for now.)
        :param limit:
        :return:
        """
        raw = '/'.join((self.rawdir, self.files['idmap']['file']))
        gu = GraphUtils(curie_map.get())
        line_counter = 0
        with open(raw, 'r', encoding="utf8") as csvfile:
            filereader = csv.reader(csvfile, delimiter='\t', quotechar='\"')
            for row in filereader:
                line_counter += 1
                (nbk_num, shortname, omim_num) = row
                gr_id = 'GeneReviews:'+nbk_num
                omim_id = 'OMIM:'+omim_num
                # add the OMIM id to the graph
                gu.addClassToGraph(self.graph, gr_id, None)
                gu.addClassToGraph(self.graph, omim_id, None)
                gu.addSynonym(self.graph, gr_id, shortname)
                gu.addSubclass(self.graph, gr_id, omim_id)

        return

    def _get_titles(self, limit):
        """
        The file processed here is of the format:
        #NBK_id GR_shortname    OMIM
        NBK1103 trimethylaminuria       136132
        NBK1103 trimethylaminuria       602079
        NBK1104 cdls    122470
        Where each of the rows represents a mapping between a gr id and an omim id.
        These are a 1:many relationship, and some of the omim ids are genes (not diseases).
        Therefore, we need to create a loose coupling here.  We make the assumption that these NBKs are
        generally higher-level grouping classes; therefore the OMIM ids are treated as subclasses.  (This
        assumption is poor for those omims that are actually genes, but we have no way of knowing what those
        are here...we will just have to deal with that for now.)
        :param limit:
        :return:
        """
        raw = '/'.join((self.rawdir, self.files['titles']['file']))
        gu = GraphUtils(curie_map.get())
        line_counter = 0
        with open(raw, 'r', encoding='latin-1') as csvfile:
            filereader = csv.reader(csvfile, delimiter='\t', quotechar='\"')
            for row in filereader:
                line_counter += 1
                (shortname, title, nbk_num) = row
                gr_id = 'GeneReviews:'+nbk_num

                gu.addClassToGraph(self.graph, gr_id, title)
                gu.addSynonym(self.graph, gr_id, shortname)

        return

    def getTestSuite(self):
        import unittest
        from tests.test_genereviews import GeneReviewsTestCase

        test_suite = unittest.TestLoader().loadTestsFromTestCase(GeneReviewsTestCase)

        return test_suite
